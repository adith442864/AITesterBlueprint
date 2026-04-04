"""
05 — FETCH JIRA → CREATE TEST PLAN (CrewAI Agent)
──────────────────────────────────────────────────
Input  : A Jira ticket ID  (e.g. VWO-48)
Output : A professional Test Plan document (.docx)

Pipeline:
  1. JiraFetcherTool   → fetches ticket from Jira REST API
  2. Test Plan Analyst  → analyzes the ticket & identifies scope
  3. Test Plan Writer   → writes the full test plan using a template
  4. Saves as .docx and opens it
"""

import os, re, sys, json, datetime, requests, webbrowser, textwrap
from pathlib import Path
from dotenv import load_dotenv
from crewai import Agent, Task, Crew, Process, LLM
from crewai.tools import tool
import litellm

load_dotenv()

# ── Rate-limit protection ─────────────────────────────────────────
os.environ["LITELLM_NUM_RETRIES"] = "5"
litellm.num_retries = 5
litellm.request_timeout = 120

# ── LLM Setup ─────────────────────────────────────────────────────
llm = LLM(
    model="groq/openai/gpt-oss-120b",
    api_key=os.getenv("GROQ_API_KEY"),
    num_retries=5,
    request_timeout=120,
)

# ══════════════════════════════════════════════════════════════════
#  JIRA CONFIGURATION
# ══════════════════════════════════════════════════════════════════
JIRA_BASE_URL  = os.getenv("JIRA_BASE_URL", "https://bugzz.atlassian.net")
JIRA_EMAIL     = os.getenv("JIRA_EMAIL")
JIRA_API_TOKEN = os.getenv("JIRA_API_TOKEN")

# ══════════════════════════════════════════════════════════════════
#  CUSTOM TOOL — Fetch Jira Ticket
# ══════════════════════════════════════════════════════════════════

@tool("JiraTicketFetcher")
def fetch_jira_ticket(ticket_id: str) -> str:
    """Fetches a Jira ticket by its ID (e.g. VWO-48) and returns
    a structured summary including title, description, priority,
    status, assignee, reporter, labels, and comments."""
    try:
        url = f"{JIRA_BASE_URL}/rest/api/3/issue/{ticket_id}"
        r = requests.get(url, auth=(JIRA_EMAIL, JIRA_API_TOKEN))
        r.raise_for_status()
        data = r.json()
        f = data["fields"]

        # ── Extract description text (Atlassian Document Format) ──
        desc_text = ""
        if f.get("description") and f["description"].get("content"):
            for block in f["description"]["content"]:
                if block.get("content"):
                    for item in block["content"]:
                        if item.get("text"):
                            desc_text += item["text"] + "\n"

        # ── Extract comments ──
        comments = ""
        if f.get("comment") and f["comment"].get("comments"):
            for c in f["comment"]["comments"][:5]:  # last 5
                author = c.get("author", {}).get("displayName", "Unknown")
                body_parts = []
                if c.get("body", {}).get("content"):
                    for block in c["body"]["content"]:
                        if block.get("content"):
                            for item in block["content"]:
                                if item.get("text"):
                                    body_parts.append(item["text"])
                comments += f"  - {author}: {' '.join(body_parts)}\n"

        # ── Build structured output ──
        result = f"""
=== JIRA TICKET: {data['key']} ===
Title       : {f.get('summary', 'N/A')}
Status      : {f.get('status', {}).get('name', 'N/A')}
Priority    : {f.get('priority', {}).get('name', 'N/A')}
Issue Type  : {f.get('issuetype', {}).get('name', 'N/A')}
Reporter    : {f.get('reporter', {}).get('displayName', 'N/A')}
Assignee    : {f.get('assignee', {}).get('displayName', 'Unassigned') if f.get('assignee') else 'Unassigned'}
Labels      : {', '.join(f.get('labels', [])) or 'None'}
Created     : {f.get('created', 'N/A')}
Updated     : {f.get('updated', 'N/A')}

Description:
{desc_text.strip() or 'No description provided.'}

Comments:
{comments.strip() or 'No comments.'}
"""
        return result.strip()

    except requests.exceptions.HTTPError as e:
        return f"Error fetching {ticket_id}: HTTP {e.response.status_code} — {e.response.text}"
    except Exception as e:
        return f"Error fetching {ticket_id}: {str(e)}"


# ══════════════════════════════════════════════════════════════════
#  TEST PLAN TEMPLATE (embedded in agent instructions)
# ══════════════════════════════════════════════════════════════════

TEST_PLAN_TEMPLATE = """
## TEST PLAN TEMPLATE — Follow this structure exactly:

### 1. TEST PLAN OVERVIEW
- **Test Plan ID**: TP-<JIRA_ID>
- **Project Name**: <from Jira ticket>
- **Feature/Module**: <from Jira ticket title>
- **Prepared By**: AI Test Plan Agent
- **Date**: <today's date>
- **Version**: 1.0

### 2. OBJECTIVE
Clearly state what this test plan aims to validate based on the Jira ticket.

### 3. SCOPE
#### 3.1 In Scope
- List all features/functionalities to be tested
#### 3.2 Out of Scope
- List items explicitly not covered

### 4. TEST STRATEGY
- **Testing Types**: (Functional, Regression, Smoke, Integration, etc.)
- **Testing Approach**: Manual / Automated / Both
- **Automation Tool**: Playwright with TypeScript (if applicable)

### 5. TEST ENVIRONMENT
- **Browser(s)**: Chrome, Firefox, Edge
- **OS**: Windows 11, macOS
- **Devices**: Desktop, Mobile (if applicable)
- **Test Data Requirements**: Describe any specific data needed

### 6. ENTRY CRITERIA
- List conditions that must be met before testing begins

### 7. EXIT CRITERIA
- List conditions that must be met for testing to be considered complete

### 8. TEST CASES

For each test case provide:
| TC ID | Test Case Title | Pre-conditions | Steps | Expected Result | Priority |
|-------|----------------|----------------|-------|-----------------|----------|

Provide at least 8-10 test cases covering:
- Positive scenarios (happy path)
- Negative scenarios (invalid inputs, boundary values)
- Edge cases
- UI/UX validations
- API validations (if applicable)

### 9. RISK ASSESSMENT
| Risk | Likelihood | Impact | Mitigation |
|------|-----------|--------|------------|

### 10. DEFECT MANAGEMENT
- **Tool**: Jira
- **Severity Levels**: P0 (Blocker), P1 (Critical), P2 (Major), P3 (Minor)
- **Bug Lifecycle**: New → In Progress → Fixed → Verified → Closed

### 11. TEST SCHEDULE
| Phase | Start Date | End Date | Owner |
|-------|-----------|----------|-------|

### 12. SIGN-OFF
- QA Lead: _______________
- Dev Lead: _______________
- Product Owner: _______________
"""


# ══════════════════════════════════════════════════════════════════
#  AGENTS
# ══════════════════════════════════════════════════════════════════

# Agent 1 — Ticket Analyst
ticket_analyst = Agent(
    role="Senior QA Analyst",
    goal="Analyze the Jira ticket thoroughly and extract all testable requirements",
    backstory="""You are a senior QA analyst with 12+ years of experience.
    You excel at reading Jira tickets and extracting:
    - Functional requirements
    - Acceptance criteria
    - Edge cases and boundary conditions
    - Non-functional requirements (performance, security, usability)
    - Dependencies and risks
    You always use the JiraTicketFetcher tool to get ticket details first.""",
    tools=[fetch_jira_ticket],
    llm=llm,
    verbose=True,
    allow_delegation=False,
)

# Agent 2 — Test Plan Writer
test_plan_writer = Agent(
    role="Test Plan Documentation Specialist",
    goal="Create a comprehensive, professional test plan document following the standard template",
    backstory=f"""You are a certified ISTQB test planning expert who writes
    detailed test plans that teams can immediately execute.
    
    You MUST follow this exact template structure:
    {TEST_PLAN_TEMPLATE}
    
    Your test plans are known for:
    - Clear, actionable test cases with step-by-step instructions
    - Thorough coverage of positive, negative, and edge cases
    - Realistic risk assessments with mitigation strategies
    - Professional formatting with proper tables and sections
    You always produce well-structured markdown output.""",
    llm=llm,
    verbose=True,
    allow_delegation=False,
)


# ══════════════════════════════════════════════════════════════════
#  ACCEPT INPUT — Jira Ticket ID
# ══════════════════════════════════════════════════════════════════

# Default to VWO-48 if no argument provided
JIRA_TICKET_ID = sys.argv[1] if len(sys.argv) > 1 else "VWO-48"
print(f"\n📡 Target Jira Ticket: {JIRA_TICKET_ID}")


# ══════════════════════════════════════════════════════════════════
#  TASKS
# ══════════════════════════════════════════════════════════════════

analysis_task = Task(
    description=f"""Fetch and analyze Jira ticket **{JIRA_TICKET_ID}** using the JiraTicketFetcher tool.

    After fetching, provide a detailed analysis:
    1. Summary of the ticket (what the feature/bug is about)
    2. List all testable requirements extracted from the ticket
    3. Identify acceptance criteria (explicit and implicit)
    4. List potential edge cases and boundary conditions
    5. Identify any risks or dependencies
    6. Suggest testing types needed (functional, regression, performance, etc.)
    
    Be thorough — the test plan writer depends on your analysis.""",
    expected_output="""A detailed analysis report containing: ticket summary,
    testable requirements, acceptance criteria, edge cases, risks,
    and recommended testing types.""",
    agent=ticket_analyst,
)

test_plan_task = Task(
    description=f"""Based on the ticket analysis, create a COMPLETE test plan
    for Jira ticket **{JIRA_TICKET_ID}**.
    
    Follow the standard test plan template EXACTLY. Include:
    - ALL 12 sections from the template
    - At least 8-10 detailed test cases in a proper table format
    - Cover: happy path, negative scenarios, edge cases, UI checks
    - Realistic risk assessment with 3-5 risks
    - Proper test schedule with phases
    - Use today's date: {datetime.date.today().strftime('%B %d, %Y')}
    
    Format everything in clean, professional markdown.
    Make sure test cases have clear, step-by-step instructions.""",
    expected_output="""A complete, professional test plan document in markdown
    format following all 12 sections of the template, with detailed test cases,
    risk assessment, and schedule.""",
    agent=test_plan_writer,
    context=[analysis_task],
)


# ══════════════════════════════════════════════════════════════════
#  RUN THE CREW
# ══════════════════════════════════════════════════════════════════

crew = Crew(
    agents=[ticket_analyst, test_plan_writer],
    tasks=[analysis_task, test_plan_task],
    process=Process.sequential,
    verbose=True,
    max_rpm=4,
)

print(f"\n🚀 Starting Test Plan Generation for {JIRA_TICKET_ID}")
print("=" * 60)

result = crew.kickoff()

analysis_output = str(analysis_task.output) if analysis_task.output else ""
test_plan_output = str(result)

print("\n" + "=" * 60)
print("📋 TEST PLAN GENERATED")
print("=" * 60)
print(test_plan_output[:500] + "...")


# ══════════════════════════════════════════════════════════════════
#  SAVE AS .DOCX
# ══════════════════════════════════════════════════════════════════

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("\n⚠️  python-docx not installed. Installing now...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE


def markdown_to_docx(md_text: str, ticket_id: str) -> str:
    """Convert the markdown test plan to a professional .docx document."""
    doc = Document()
    
    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Custom styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ── Title page ──
    doc.add_paragraph()  # spacer
    title = doc.add_heading(f'Test Plan — {ticket_id}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        run.font.size = Pt(28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Generated by AI Test Plan Agent\n{datetime.date.today().strftime("%B %d, %Y")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()  # spacer

    # ── Parse markdown into document ──
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # ── Headers ──
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        elif stripped.startswith('## '):
            h = doc.add_heading(stripped[3:], level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x0d, 0x3b, 0x8f)
        elif stripped.startswith('# '):
            h = doc.add_heading(stripped[2:], level=0)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x0d, 0x3b, 0x8f)

        # ── Tables (markdown pipe tables) ──
        elif stripped.startswith('|') and stripped.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # back up one since we'll increment at the end

            # Parse table
            if len(table_lines) >= 2:
                # Header row
                headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                
                # Skip separator row (---|---|...)
                data_start = 1
                if len(table_lines) > 1 and all(c.strip().replace('-', '').replace(':', '') == '' for c in table_lines[1].split('|')[1:-1]):
                    data_start = 2

                # Data rows
                data_rows = []
                for tl in table_lines[data_start:]:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    data_rows.append(cells)

                # Create Word table
                num_cols = len(headers)
                table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Header row styling
                header_row = table.rows[0]
                for j, h_text in enumerate(headers):
                    if j < num_cols:
                        cell = header_row.cells[j]
                        cell.text = h_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        from docx.oxml.ns import qn
                        shading = cell._element.get_or_add_tcPr()
                        shading_elm = shading.makeelement(qn('w:shd'), {
                            qn('w:fill'): '1a56db',
                            qn('w:val'): 'clear'
                        })
                        shading.append(shading_elm)

                # Data rows
                for row_idx, row_data in enumerate(data_rows):
                    row = table.rows[row_idx + 1]
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = row.cells[col_idx]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                doc.add_paragraph()  # spacer after table

        # ── Horizontal rules ──
        elif stripped.startswith('---') or stripped.startswith('___'):
            doc.add_paragraph('─' * 60)

        # ── Bullet points ──
        elif stripped.startswith('- **') or stripped.startswith('* **'):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold within bullet
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(11)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(11)

        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold/inline formatting
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # ── Numbered items ──
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # ── Regular paragraph ──
        else:
            p = doc.add_paragraph()
            # Handle bold
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f'Generated by AI Test Plan Agent — CrewAI × Groq × Jira\n{datetime.datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Save ──
    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)
    filename = f"TestPlan_{ticket_id}_{datetime.date.today().strftime('%Y%m%d')}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))
    return str(filepath)


# ── Generate the DOCX ─────────────────────────────────────────────
print("\n📝 Converting test plan to DOCX ...")
docx_path = markdown_to_docx(test_plan_output, JIRA_TICKET_ID)
print(f"✅ Test Plan saved to: {docx_path}")

# ── Open the document ─────────────────────────────────────────────
webbrowser.open(f"file://{Path(docx_path).resolve()}")
print("🌐 Opening document ...")
print("\n🎉 Done! Test plan generation complete.")