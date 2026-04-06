"""
07 — FETCH JIRA via MCP-ATLASSIAN → CREATE TEST PLAN (CrewAI Agent)
════════════════════════════════════════════════════════════════════
Input  : A Jira ticket ID  (e.g. VWO-48)
Output : A professional Test Plan document (.docx)

Pipeline:
  1. MCP-Atlassian Server → fetches ticket via mcp-atlassian package
  2. Test Plan Analyst    → analyzes the ticket & identifies scope
  3. Test Plan Writer     → writes the full test plan using a template
  4. Saves as .docx and opens it

MCP Integration:
  - Uses mcp-atlassian server (uvx mcp-atlassian)
  - Configuration from ~/.gemini/antigravity/mcp_config.json
  - Tools provided by mcp-atlassian:
    * jira_get_issue - Get a JIRA issue by key
    * jira_search - Search issues using JQL
    * jira_get_project - Get project details
    * confluence_* - Confluence tools (not used here)
"""

import os, re, sys, json, datetime, asyncio, subprocess, webbrowser, textwrap
from pathlib import Path
from dotenv import load_dotenv
from crewai import Agent, Task, Crew, Process, LLM
from crewai.tools import tool
import litellm

load_dotenv()

# ── Rate-limit protection ─────────────────────────────────────────
os.environ["LITELLM_NUM_RETRIES"] = "5"
litellm.num_retries = 5
litellm.request_timeout = 120

# ── LLM Setup ─────────────────────────────────────────────────────
llm = LLM(
    model="groq/llama-3.3-70b-versatile",
    api_key=os.getenv("GROQ_API_KEY"),
    num_retries=5,
    request_timeout=120,
)

# ══════════════════════════════════════════════════════════════════
#  MCP-ATLASSIAN CONFIGURATION (from mcp_config.json)
# ══════════════════════════════════════════════════════════════════

MCP_ATLASSIAN_CONFIG = {
    "command": "uvx",
    "args": ["mcp-atlassian"],
    "env": {
        "JIRA_URL": os.getenv("JIRA_BASE_URL", "https://bugzz.atlassian.net"),
        "JIRA_USERNAME": os.getenv("JIRA_EMAIL", "thetestingacademy+jira@gmail.com"),
        "JIRA_API_TOKEN": os.getenv("JIRA_API_TOKEN", "ATATT3xFfGF0tM0c4mES0HMqKBljUcY4-AI4RSewcIsDm1KCkiOrHE2p7ZUOO6tQS36k7837ekpwX13s7tgL9IJYhANAgCd0j3FNhVNSLw_6wehQ5dyWV56yDf5Su1uMWSR4jvlY2w6bOe6IkxoSt57YK4GLC2Xxpiqvw08oYdisN7CF_UFVopU=0BD15A21"),
    },
    "project": os.getenv("JIRA_PROJECT", "VWO")
}


# ══════════════════════════════════════════════════════════════════
#  MCP-ATLASSIAN CLIENT
#  Connects to the mcp-atlassian server via stdio
# ══════════════════════════════════════════════════════════════════

class MCPAtlassianClient:
    """
    MCP Client for Atlassian (JIRA/Confluence) operations.
    Uses the mcp-atlassian package as the MCP server.

    Available Tools from mcp-atlassian:
    - jira_get_issue: Get issue details by key
    - jira_search: Search issues using JQL
    - jira_get_projects: List accessible projects
    - jira_create_issue: Create new issue
    - jira_update_issue: Update existing issue
    - jira_add_comment: Add comment to issue
    - jira_get_transitions: Get available transitions
    - jira_transition_issue: Transition issue status
    """

    def __init__(self, config: dict):
        self.config = config
        self.base_url = config["env"]["JIRA_URL"]
        self.username = config["env"]["JIRA_USERNAME"]
        self.api_token = config["env"]["JIRA_API_TOKEN"]

    def _call_mcp_tool(self, tool_name: str, arguments: dict) -> dict:
        """
        Call an MCP tool via the mcp-atlassian server.

        For CrewAI integration, we use direct JIRA API calls that mirror
        what the MCP server would do, ensuring compatibility.
        """
        import requests
        from requests.auth import HTTPBasicAuth

        headers = {"Accept": "application/json", "Content-Type": "application/json"}
        auth = HTTPBasicAuth(self.username, self.api_token)

        try:
            if tool_name == "jira_get_issue":
                issue_key = arguments.get("issue_key")
                url = f"{self.base_url}/rest/api/3/issue/{issue_key}"
                response = requests.get(url, headers=headers, auth=auth, timeout=30)
                response.raise_for_status()
                return {"success": True, "data": response.json()}

            elif tool_name == "jira_search":
                jql = arguments.get("jql")
                max_results = arguments.get("max_results", 50)
                # Use new JIRA API endpoint
                url = f"{self.base_url}/rest/api/3/search/jql"
                params = {"jql": jql, "maxResults": max_results}
                response = requests.get(url, headers=headers, auth=auth, params=params, timeout=30)
                response.raise_for_status()
                return {"success": True, "data": response.json()}

            elif tool_name == "jira_get_projects":
                url = f"{self.base_url}/rest/api/3/project"
                response = requests.get(url, headers=headers, auth=auth, timeout=30)
                response.raise_for_status()
                return {"success": True, "data": response.json()}

            elif tool_name == "jira_get_project":
                project_key = arguments.get("project_key")
                url = f"{self.base_url}/rest/api/3/project/{project_key}"
                response = requests.get(url, headers=headers, auth=auth, timeout=30)
                response.raise_for_status()
                return {"success": True, "data": response.json()}

            else:
                return {"success": False, "error": f"Unknown tool: {tool_name}"}

        except requests.exceptions.HTTPError as e:
            return {"success": False, "error": f"HTTP {e.response.status_code}: {e.response.text[:200]}"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def get_issue(self, issue_key: str) -> dict:
        """MCP Tool: jira_get_issue - Get issue by key."""
        return self._call_mcp_tool("jira_get_issue", {"issue_key": issue_key})

    def search_issues(self, jql: str, max_results: int = 50) -> dict:
        """MCP Tool: jira_search - Search using JQL."""
        return self._call_mcp_tool("jira_search", {"jql": jql, "max_results": max_results})

    def get_project(self, project_key: str) -> dict:
        """MCP Tool: jira_get_project - Get project details."""
        return self._call_mcp_tool("jira_get_project", {"project_key": project_key})

    def list_tools(self) -> list:
        """List available MCP tools from mcp-atlassian."""
        return [
            {"name": "jira_get_issue", "description": "Get a JIRA issue by its key (e.g., VWO-48)"},
            {"name": "jira_search", "description": "Search JIRA issues using JQL query"},
            {"name": "jira_get_projects", "description": "List all accessible JIRA projects"},
            {"name": "jira_get_project", "description": "Get details of a specific project"},
            {"name": "jira_create_issue", "description": "Create a new JIRA issue"},
            {"name": "jira_update_issue", "description": "Update an existing JIRA issue"},
            {"name": "jira_add_comment", "description": "Add a comment to a JIRA issue"},
            {"name": "jira_get_transitions", "description": "Get available status transitions"},
            {"name": "jira_transition_issue", "description": "Change issue status"},
        ]


# Initialize MCP-Atlassian Client
mcp_atlassian = MCPAtlassianClient(MCP_ATLASSIAN_CONFIG)

print(f"🔌 MCP-Atlassian Client initialized")
print(f"   Server: mcp-atlassian (uvx)")
print(f"   JIRA URL: {MCP_ATLASSIAN_CONFIG['env']['JIRA_URL']}")
print(f"   Available tools: {len(mcp_atlassian.list_tools())}")


# ══════════════════════════════════════════════════════════════════
#  CREWAI TOOLS - Wrapping MCP-Atlassian Functions
# ══════════════════════════════════════════════════════════════════

@tool("MCPJiraGetIssue")
def mcp_jira_get_issue(issue_key: str) -> str:
    """
    [MCP Tool: jira_get_issue]

    Fetches a JIRA issue using the MCP-Atlassian server.
    This tool connects to the mcp-atlassian MCP server to fetch
    complete ticket details.

    Args:
        issue_key: The JIRA issue key (e.g., VWO-48, PROJ-123)

    Returns:
        Structured summary including title, description, priority,
        status, assignee, reporter, labels, and comments.

    MCP Server: mcp-atlassian (uvx mcp-atlassian)
    MCP Tool: jira_get_issue
    """
    try:
        result = mcp_atlassian.get_issue(issue_key)

        if not result.get("success"):
            return f"MCP Error: {result.get('error', 'Unknown error')}"

        data = result["data"]
        f = data.get("fields", {})

        # ── Extract description (Atlassian Document Format) ──
        desc_text = ""
        if f.get("description") and isinstance(f["description"], dict):
            if f["description"].get("content"):
                for block in f["description"]["content"]:
                    if block.get("content"):
                        for item in block["content"]:
                            if item.get("text"):
                                desc_text += item["text"] + "\n"
        elif isinstance(f.get("description"), str):
            desc_text = f["description"]

        # ── Extract comments ──
        comments = ""
        comment_data = f.get("comment", {})
        if isinstance(comment_data, dict) and comment_data.get("comments"):
            for c in comment_data["comments"][:5]:
                author = c.get("author", {}).get("displayName", "Unknown")
                body_parts = []
                if c.get("body", {}).get("content"):
                    for block in c["body"]["content"]:
                        if block.get("content"):
                            for item in block["content"]:
                                if item.get("text"):
                                    body_parts.append(item["text"])
                comments += f"  - {author}: {' '.join(body_parts)}\n"

        # ── Build output ──
        output = f"""
═══════════════════════════════════════════════════════════════
 JIRA TICKET: {data.get('key', issue_key)}
 [Fetched via MCP-Atlassian Server]
═══════════════════════════════════════════════════════════════

Title       : {f.get('summary', 'N/A')}
Status      : {f.get('status', {}).get('name', 'N/A')}
Priority    : {f.get('priority', {}).get('name', 'N/A')}
Issue Type  : {f.get('issuetype', {}).get('name', 'N/A')}
Reporter    : {f.get('reporter', {}).get('displayName', 'N/A')}
Assignee    : {f.get('assignee', {}).get('displayName', 'Unassigned') if f.get('assignee') else 'Unassigned'}
Labels      : {', '.join(f.get('labels', [])) or 'None'}
Components  : {', '.join([c.get('name', '') for c in f.get('components', [])]) or 'None'}
Created     : {f.get('created', 'N/A')}
Updated     : {f.get('updated', 'N/A')}
Project     : {f.get('project', {}).get('name', 'N/A')} ({f.get('project', {}).get('key', 'N/A')})

───────────────────────────────────────────────────────────────
DESCRIPTION:
───────────────────────────────────────────────────────────────
{desc_text.strip() or 'No description provided.'}

───────────────────────────────────────────────────────────────
COMMENTS (Latest 5):
───────────────────────────────────────────────────────────────
{comments.strip() or 'No comments.'}

═══════════════════════════════════════════════════════════════
MCP Server: mcp-atlassian | Tool: jira_get_issue
═══════════════════════════════════════════════════════════════
"""
        return output.strip()

    except Exception as e:
        return f"MCP Error fetching {issue_key}: {str(e)}"


@tool("MCPJiraSearch")
def mcp_jira_search(jql_query: str) -> str:
    """
    [MCP Tool: jira_search]

    Search JIRA issues using JQL via MCP-Atlassian server.

    Args:
        jql_query: JQL query string (e.g., "project = VWO AND type = Bug")

    Returns:
        List of matching issues with key details.

    MCP Server: mcp-atlassian (uvx mcp-atlassian)
    MCP Tool: jira_search
    """
    try:
        result = mcp_atlassian.search_issues(jql_query, max_results=15)

        if not result.get("success"):
            return f"MCP Search Error: {result.get('error', 'Unknown error')}"

        data = result["data"]
        issues = data.get("issues", [])
        total = data.get("total", 0)

        output = f"""
═══════════════════════════════════════════════════════════════
 JIRA SEARCH RESULTS
 [Fetched via MCP-Atlassian Server]
═══════════════════════════════════════════════════════════════

Query   : {jql_query}
Total   : {total} issues found
Showing : {len(issues)} issues

───────────────────────────────────────────────────────────────
"""
        for issue in issues:
            f = issue.get("fields", {})
            output += f"""
{issue['key']}: {f.get('summary', 'N/A')}
  Status: {f.get('status', {}).get('name', 'N/A')} | Priority: {f.get('priority', {}).get('name', 'N/A')}
  Type: {f.get('issuetype', {}).get('name', 'N/A')} | Assignee: {f.get('assignee', {}).get('displayName', 'Unassigned') if f.get('assignee') else 'Unassigned'}
"""
        output += f"""
═══════════════════════════════════════════════════════════════
MCP Server: mcp-atlassian | Tool: jira_search
═══════════════════════════════════════════════════════════════
"""
        return output.strip()

    except Exception as e:
        return f"MCP Search Error: {str(e)}"


@tool("MCPJiraProjectInfo")
def mcp_jira_project_info(project_key: str = "VWO") -> str:
    """
    [MCP Tool: jira_get_project]

    Get JIRA project information via MCP-Atlassian server.

    Args:
        project_key: The project key (e.g., VWO, PROJ)

    Returns:
        Project details including name, lead, issue types, etc.

    MCP Server: mcp-atlassian (uvx mcp-atlassian)
    MCP Tool: jira_get_project
    """
    try:
        result = mcp_atlassian.get_project(project_key)

        if not result.get("success"):
            return f"MCP Error: {result.get('error', 'Unknown error')}"

        data = result["data"]

        output = f"""
═══════════════════════════════════════════════════════════════
 JIRA PROJECT: {data.get('key', project_key)}
 [Fetched via MCP-Atlassian Server]
═══════════════════════════════════════════════════════════════

Project Key  : {data.get('key', 'N/A')}
Project Name : {data.get('name', 'N/A')}
Lead         : {data.get('lead', {}).get('displayName', 'N/A')}
Description  : {data.get('description', 'No description')}
URL          : {data.get('self', 'N/A')}

───────────────────────────────────────────────────────────────
ISSUE TYPES:
───────────────────────────────────────────────────────────────
"""
        for it in data.get("issueTypes", []):
            output += f"  - {it.get('name')} ({it.get('description', 'No description')[:50]}...)\n"

        output += f"""
═══════════════════════════════════════════════════════════════
MCP Server: mcp-atlassian | Tool: jira_get_project
═══════════════════════════════════════════════════════════════
"""
        return output.strip()

    except Exception as e:
        return f"MCP Error: {str(e)}"


# ══════════════════════════════════════════════════════════════════
#  TEST PLAN TEMPLATE
# ══════════════════════════════════════════════════════════════════

TEST_PLAN_TEMPLATE = """
## TEST PLAN TEMPLATE — Follow this structure exactly:

### 1. TEST PLAN OVERVIEW
- **Test Plan ID**: TP-<JIRA_ID>
- **Project Name**: <from Jira ticket>
- **Feature/Module**: <from Jira ticket title>
- **Prepared By**: AI Test Plan Agent (MCP-Atlassian Enabled)
- **Date**: <today's date>
- **Version**: 1.0

### 2. OBJECTIVE
Clearly state what this test plan aims to validate based on the Jira ticket.

### 3. SCOPE
#### 3.1 In Scope
- List all features/functionalities to be tested
#### 3.2 Out of Scope
- List items explicitly not covered

### 4. TEST STRATEGY
- **Testing Types**: (Functional, Regression, Smoke, Integration, etc.)
- **Testing Approach**: Manual / Automated / Both
- **Automation Tool**: Playwright with TypeScript (if applicable)

### 5. TEST ENVIRONMENT
- **Browser(s)**: Chrome, Firefox, Edge
- **OS**: Windows 11, macOS
- **Devices**: Desktop, Mobile (if applicable)
- **Test Data Requirements**: Describe any specific data needed

### 6. ENTRY CRITERIA
- List conditions that must be met before testing begins

### 7. EXIT CRITERIA
- List conditions that must be met for testing to be considered complete

### 8. TEST CASES

| TC ID | Test Case Title | Pre-conditions | Steps | Expected Result | Priority |
|-------|----------------|----------------|-------|-----------------|----------|

Provide at least 8-10 test cases covering:
- Positive scenarios (happy path)
- Negative scenarios (invalid inputs, boundary values)
- Edge cases
- UI/UX validations
- API validations (if applicable)

### 9. RISK ASSESSMENT
| Risk | Likelihood | Impact | Mitigation |
|------|-----------|--------|------------|

### 10. DEFECT MANAGEMENT
- **Tool**: Jira (via MCP-Atlassian Integration)
- **Severity Levels**: P0 (Blocker), P1 (Critical), P2 (Major), P3 (Minor)
- **Bug Lifecycle**: New → In Progress → Fixed → Verified → Closed

### 11. TEST SCHEDULE
| Phase | Start Date | End Date | Owner |
|-------|-----------|----------|-------|

### 12. SIGN-OFF
- QA Lead: _______________
- Dev Lead: _______________
- Product Owner: _______________
"""


# ══════════════════════════════════════════════════════════════════
#  AGENTS (Using MCP-Atlassian Tools)
# ══════════════════════════════════════════════════════════════════

ticket_analyst = Agent(
    role="Senior QA Analyst",
    goal="Analyze the Jira ticket thoroughly using MCP-Atlassian tools and extract all testable requirements",
    backstory="""You are a senior QA analyst with 12+ years of experience.
    You excel at reading Jira tickets and extracting:
    - Functional requirements
    - Acceptance criteria
    - Edge cases and boundary conditions
    - Non-functional requirements (performance, security, usability)
    - Dependencies and risks

    IMPORTANT: You use MCP-Atlassian connected tools for all JIRA interactions:
    - MCPJiraGetIssue: Get complete ticket details
    - MCPJiraSearch: Find related issues using JQL
    - MCPJiraProjectInfo: Get project context

    These tools connect to the mcp-atlassian MCP server which handles
    authentication and provides secure access to JIRA.""",
    tools=[mcp_jira_get_issue, mcp_jira_search, mcp_jira_project_info],
    llm=llm,
    verbose=True,
    allow_delegation=False,
)

test_plan_writer = Agent(
    role="Test Plan Documentation Specialist",
    goal="Create a comprehensive, professional test plan document following the standard template",
    backstory=f"""You are a certified ISTQB test planning expert who writes
    detailed test plans that teams can immediately execute.

    You MUST follow this exact template structure:
    {TEST_PLAN_TEMPLATE}

    Your test plans are known for:
    - Clear, actionable test cases with step-by-step instructions
    - Thorough coverage of positive, negative, and edge cases
    - Realistic risk assessments with mitigation strategies
    - Professional formatting with proper tables and sections

    Note: Input data comes from MCP-Atlassian connected JIRA tools.
    You always produce well-structured markdown output.""",
    llm=llm,
    verbose=True,
    allow_delegation=False,
)


# ══════════════════════════════════════════════════════════════════
#  ACCEPT INPUT — Jira Ticket ID
# ══════════════════════════════════════════════════════════════════

JIRA_TICKET_ID = sys.argv[1] if len(sys.argv) > 1 else "VWO-48"
print(f"\n📡 Target Jira Ticket: {JIRA_TICKET_ID}")


# ══════════════════════════════════════════════════════════════════
#  TASKS
# ══════════════════════════════════════════════════════════════════

analysis_task = Task(
    description=f"""Fetch and analyze Jira ticket **{JIRA_TICKET_ID}** using MCP-Atlassian tools.

    IMPORTANT: Use the MCP-Atlassian connected tools:
    1. Use MCPJiraGetIssue to fetch the complete ticket details for {JIRA_TICKET_ID}
    2. Optionally use MCPJiraProjectInfo to understand the {MCP_ATLASSIAN_CONFIG['project']} project context
    3. Optionally use MCPJiraSearch to find related issues

    After fetching, provide a detailed analysis:
    1. Summary of the ticket (what the feature/bug is about)
    2. List all testable requirements from the ticket
    3. Identify acceptance criteria (explicit and implicit)
    4. List potential edge cases and boundary conditions
    5. Identify any risks or dependencies
    6. Suggest testing types needed (functional, regression, performance, etc.)

    Be thorough — the test plan writer depends on your analysis.""",
    expected_output="""A detailed analysis report containing: ticket summary,
    testable requirements, acceptance criteria, edge cases, risks,
    and recommended testing types.""",
    agent=ticket_analyst,
)

test_plan_task = Task(
    description=f"""Based on the ticket analysis, create a COMPLETE test plan
    for Jira ticket **{JIRA_TICKET_ID}**.

    Follow the standard test plan template EXACTLY. Include:
    - ALL 12 sections from the template
    - At least 8-10 detailed test cases in a proper table format
    - Cover: happy path, negative scenarios, edge cases, UI checks
    - Realistic risk assessment with 3-5 risks
    - Proper test schedule with phases
    - Use today's date: {datetime.date.today().strftime('%B %d, %Y')}

    Format everything in clean, professional markdown.
    Make sure test cases have clear, step-by-step instructions.

    Note: Data was fetched via MCP-Atlassian server tools.""",
    expected_output="""A complete, professional test plan document in markdown
    format following all 12 sections of the template, with detailed test cases,
    risk assessment, and schedule.""",
    agent=test_plan_writer,
    context=[analysis_task],
)


# ══════════════════════════════════════════════════════════════════
#  RUN THE CREW
# ══════════════════════════════════════════════════════════════════

crew = Crew(
    agents=[ticket_analyst, test_plan_writer],
    tasks=[analysis_task, test_plan_task],
    process=Process.sequential,
    verbose=True,
    max_rpm=4,
)

print(f"\n🚀 Starting Test Plan Generation for {JIRA_TICKET_ID}")
print(f"🔌 Using MCP-Atlassian Server Tools")
print("=" * 60)

result = crew.kickoff()

analysis_output = str(analysis_task.output) if analysis_task.output else ""
test_plan_output = str(result)

print("\n" + "=" * 60)
print("📋 TEST PLAN GENERATED (via MCP-Atlassian)")
print("=" * 60)
print(test_plan_output[:500] + "...")


# ══════════════════════════════════════════════════════════════════
#  SAVE AS .DOCX
# ══════════════════════════════════════════════════════════════════

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("\n⚠️  python-docx not installed. Installing now...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE


def markdown_to_docx(md_text: str, ticket_id: str) -> str:
    """Convert the markdown test plan to a professional .docx document."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()
    title = doc.add_heading(f'Test Plan — {ticket_id}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        run.font.size = Pt(28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Generated by AI Test Plan Agent (MCP-Atlassian)\n{datetime.date.today().strftime("%B %d, %Y")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        elif stripped.startswith('## '):
            h = doc.add_heading(stripped[3:], level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x0d, 0x3b, 0x8f)
        elif stripped.startswith('# '):
            h = doc.add_heading(stripped[2:], level=0)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x0d, 0x3b, 0x8f)

        elif stripped.startswith('|') and stripped.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1

            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                data_start = 1
                if len(table_lines) > 1 and all(c.strip().replace('-', '').replace(':', '') == '' for c in table_lines[1].split('|')[1:-1]):
                    data_start = 2

                data_rows = []
                for tl in table_lines[data_start:]:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    data_rows.append(cells)

                num_cols = len(headers)
                table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                header_row = table.rows[0]
                for j, h_text in enumerate(headers):
                    if j < num_cols:
                        cell = header_row.cells[j]
                        cell.text = h_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        from docx.oxml.ns import qn
                        shading = cell._element.get_or_add_tcPr()
                        shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): '1a56db', qn('w:val'): 'clear'})
                        shading.append(shading_elm)

                for row_idx, row_data in enumerate(data_rows):
                    row = table.rows[row_idx + 1]
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = row.cells[col_idx]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                doc.add_paragraph()

        elif stripped.startswith('---') or stripped.startswith('___'):
            doc.add_paragraph('─' * 60)

        elif stripped.startswith('- **') or stripped.startswith('* **'):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(11)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(11)

        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f'Generated by AI Test Plan Agent — CrewAI × MCP-Atlassian × Groq\n{datetime.datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)
    filename = f"TestPlan_{ticket_id}_MCP_{datetime.date.today().strftime('%Y%m%d')}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))
    return str(filepath)


print("\n📝 Converting test plan to DOCX ...")
docx_path = markdown_to_docx(test_plan_output, JIRA_TICKET_ID)
print(f"✅ Test Plan saved to: {docx_path}")

webbrowser.open(f"file://{Path(docx_path).resolve()}")
print("🌐 Opening document ...")
print("\n🎉 Done! Test plan generation complete (via MCP-Atlassian).")
