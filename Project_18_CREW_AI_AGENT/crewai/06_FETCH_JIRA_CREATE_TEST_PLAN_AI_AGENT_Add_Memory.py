"""
06 — FETCH JIRA → CREATE TEST PLAN (with CrewAI Memory)
────────────────────────────────────────────────────────
Same pipeline as 05, but with ALL THREE memory types enabled:

  ┌─────────────────────────────────────────────────────────────┐
  │  MEMORY TYPE       │  WHAT IT DOES                         │
  ├────────────────────┼───────────────────────────────────────┤
  │  Short-Term Memory │  Remembers context WITHIN this run.   │
  │                    │  The Analyst's output stays fresh for  │
  │                    │  the Writer — no context loss.         │
  ├────────────────────┼───────────────────────────────────────┤
  │  Long-Term Memory  │  Persists learnings ACROSS runs.      │
  │                    │  Run VWO-48 today, VWO-49 tomorrow —  │
  │                    │  test plans get better each time.      │
  ├────────────────────┼───────────────────────────────────────┤
  │  Entity Memory     │  Tracks specific ENTITIES (ticket IDs,│
  │                    │  people, modules like "Shopping Cart").│
  │                    │  Builds a knowledge graph over time.   │
  └─────────────────────────────────────────────────────────────┘

How to enable in CrewAI:
  - Set  memory=True       on the Crew  (enables short-term + long-term)
  - Set  memory=True       + entity tracking is automatic
  - Optionally set  embedder  config for the vector store used by memory

Input  : Jira ticket ID (e.g. VWO-48)
Output : Professional Test Plan (.docx)
"""

import os, re, sys, json, datetime, requests, webbrowser, textwrap
from pathlib import Path
from dotenv import load_dotenv
from crewai import Agent, Task, Crew, Process, LLM
from crewai.tools import tool
import litellm

load_dotenv()

# ── Rate-limit protection ─────────────────────────────────────────
os.environ["LITELLM_NUM_RETRIES"] = "5"
litellm.num_retries = 5
litellm.request_timeout = 120

# ── LLM Setup ─────────────────────────────────────────────────────
llm = LLM(
    model="groq/openai/gpt-oss-120b",
    api_key=os.getenv("GROQ_API_KEY"),
    num_retries=5,
    request_timeout=120,
)

# ══════════════════════════════════════════════════════════════════
#  JIRA CONFIGURATION
# ══════════════════════════════════════════════════════════════════
JIRA_BASE_URL  = os.getenv("JIRA_BASE_URL", "https://bugzz.atlassian.net")
JIRA_EMAIL     = os.getenv("JIRA_EMAIL")
JIRA_API_TOKEN = os.getenv("JIRA_API_TOKEN")

# ══════════════════════════════════════════════════════════════════
#  CUSTOM TOOL — Fetch Jira Ticket
# ══════════════════════════════════════════════════════════════════

@tool("JiraTicketFetcher")
def fetch_jira_ticket(ticket_id: str) -> str:
    """Fetches a Jira ticket by its ID (e.g. VWO-48) and returns
    a structured summary including title, description, priority,
    status, assignee, reporter, labels, and comments."""
    try:
        url = f"{JIRA_BASE_URL}/rest/api/3/issue/{ticket_id}"
        r = requests.get(url, auth=(JIRA_EMAIL, JIRA_API_TOKEN))
        r.raise_for_status()
        data = r.json()
        f = data["fields"]

        # Extract description text (Atlassian Document Format)
        desc_text = ""
        if f.get("description") and f["description"].get("content"):
            for block in f["description"]["content"]:
                if block.get("content"):
                    for item in block["content"]:
                        if item.get("text"):
                            desc_text += item["text"] + "\n"

        # Extract comments
        comments = ""
        if f.get("comment") and f["comment"].get("comments"):
            for c in f["comment"]["comments"][:5]:
                author = c.get("author", {}).get("displayName", "Unknown")
                body_parts = []
                if c.get("body", {}).get("content"):
                    for block in c["body"]["content"]:
                        if block.get("content"):
                            for item in block["content"]:
                                if item.get("text"):
                                    body_parts.append(item["text"])
                comments += f"  - {author}: {' '.join(body_parts)}\n"

        result = f"""
=== JIRA TICKET: {data['key']} ===
Title       : {f.get('summary', 'N/A')}
Status      : {f.get('status', {}).get('name', 'N/A')}
Priority    : {f.get('priority', {}).get('name', 'N/A')}
Issue Type  : {f.get('issuetype', {}).get('name', 'N/A')}
Reporter    : {f.get('reporter', {}).get('displayName', 'N/A')}
Assignee    : {f.get('assignee', {}).get('displayName', 'Unassigned') if f.get('assignee') else 'Unassigned'}
Labels      : {', '.join(f.get('labels', [])) or 'None'}
Created     : {f.get('created', 'N/A')}
Updated     : {f.get('updated', 'N/A')}

Description:
{desc_text.strip() or 'No description provided.'}

Comments:
{comments.strip() or 'No comments.'}
"""
        return result.strip()

    except requests.exceptions.HTTPError as e:
        return f"Error fetching {ticket_id}: HTTP {e.response.status_code} — {e.response.text}"
    except Exception as e:
        return f"Error fetching {ticket_id}: {str(e)}"


# ══════════════════════════════════════════════════════════════════
#  TEST PLAN TEMPLATE
# ══════════════════════════════════════════════════════════════════

TEST_PLAN_TEMPLATE = """
## TEST PLAN TEMPLATE — Follow this structure exactly:

### 1. TEST PLAN OVERVIEW
- **Test Plan ID**: TP-<JIRA_ID>
- **Project Name**: <from Jira ticket>
- **Feature/Module**: <from Jira ticket title>
- **Prepared By**: AI Test Plan Agent
- **Date**: <today's date>
- **Version**: 1.0

### 2. OBJECTIVE
Clearly state what this test plan aims to validate.

### 3. SCOPE
#### 3.1 In Scope
- List all features/functionalities to be tested
#### 3.2 Out of Scope
- List items explicitly not covered

### 4. TEST STRATEGY
- **Testing Types**: (Functional, Regression, Smoke, Integration, etc.)
- **Testing Approach**: Manual / Automated / Both
- **Automation Tool**: Playwright with TypeScript (if applicable)

### 5. TEST ENVIRONMENT
- **Browser(s)**: Chrome, Firefox, Edge
- **OS**: Windows 11, macOS
- **Devices**: Desktop, Mobile (if applicable)
- **Test Data Requirements**: Describe any specific data needed

### 6. ENTRY CRITERIA
- Conditions that must be met before testing begins

### 7. EXIT CRITERIA
- Conditions that must be met for testing to be complete

### 8. TEST CASES
| TC ID | Test Case Title | Pre-conditions | Steps | Expected Result | Priority |
Provide at least 8-10 test cases covering positive, negative, and edge cases.

### 9. RISK ASSESSMENT
| Risk | Likelihood | Impact | Mitigation |

### 10. DEFECT MANAGEMENT
- Tool: Jira
- Severity Levels: P0-P3
- Bug Lifecycle: New → In Progress → Fixed → Verified → Closed

### 11. TEST SCHEDULE
| Phase | Start Date | End Date | Owner |

### 12. SIGN-OFF
- QA Lead / Dev Lead / Product Owner
"""


# ══════════════════════════════════════════════════════════════════
#  AGENTS
# ══════════════════════════════════════════════════════════════════

ticket_analyst = Agent(
    role="Senior QA Analyst",
    goal="Analyze the Jira ticket thoroughly and extract all testable requirements",
    backstory="""You are a senior QA analyst with 12+ years of experience.
    You excel at reading Jira tickets and extracting:
    - Functional requirements and acceptance criteria
    - Edge cases and boundary conditions
    - Non-functional requirements (performance, security, usability)
    - Dependencies and risks
    
    MEMORY ADVANTAGE: You remember patterns from previous ticket analyses.
    If you've seen similar bugs or modules before, leverage that knowledge
    to provide richer analysis. You always use the JiraTicketFetcher tool first.""",
    tools=[fetch_jira_ticket],
    llm=llm,
    verbose=True,
    allow_delegation=False,
    # ─── Agent-level memory config ───
    # When Crew memory=True, this agent automatically gets:
    #   • Short-term: remembers its own reasoning within this execution
    #   • Long-term: recalls insights from past executions
    #   • Entity: recognizes entities (ticket IDs, modules, reporters)
)

test_plan_writer = Agent(
    role="Test Plan Documentation Specialist",
    goal="Create a comprehensive, professional test plan following the standard template",
    backstory=f"""You are a certified ISTQB test planning expert who writes
    detailed test plans that teams can immediately execute.
    
    You MUST follow this exact template:
    {TEST_PLAN_TEMPLATE}
    
    MEMORY ADVANTAGE: You recall test plan patterns from previous runs.
    If you wrote a test plan for a similar module before, you reuse
    proven test case patterns and risk assessments. Your test plans
    improve with each execution.
    
    Your output is always clean, professional markdown.""",
    llm=llm,
    verbose=True,
    allow_delegation=False,
)


# ══════════════════════════════════════════════════════════════════
#  INPUT — Jira Ticket ID
# ══════════════════════════════════════════════════════════════════

JIRA_TICKET_ID = sys.argv[1] if len(sys.argv) > 1 else "VWO-48"
print(f"\n📡 Target Jira Ticket: {JIRA_TICKET_ID}")


# ══════════════════════════════════════════════════════════════════
#  TASKS
# ══════════════════════════════════════════════════════════════════

analysis_task = Task(
    description=f"""Fetch and analyze Jira ticket **{JIRA_TICKET_ID}** using the JiraTicketFetcher tool.

    After fetching, provide a detailed analysis:
    1. Summary of the ticket (what the feature/bug is about)
    2. List all testable requirements extracted from the ticket
    3. Identify acceptance criteria (explicit and implicit)
    4. List potential edge cases and boundary conditions
    5. Identify any risks or dependencies
    6. Suggest testing types needed (functional, regression, performance, etc.)
    
    IMPORTANT: If you have memory of analyzing similar tickets or modules
    in past executions, leverage that knowledge to enrich your analysis.""",
    expected_output="""A detailed analysis report with: ticket summary,
    testable requirements, acceptance criteria, edge cases, risks,
    and recommended testing types.""",
    agent=ticket_analyst,
)

test_plan_task = Task(
    description=f"""Based on the ticket analysis, create a COMPLETE test plan
    for Jira ticket **{JIRA_TICKET_ID}**.
    
    Follow the standard test plan template EXACTLY. Include ALL 12 sections:
    - At least 8-10 detailed test cases in table format
    - Cover: happy path, negative scenarios, edge cases, UI checks
    - Realistic risk assessment with 3-5 risks
    - Use today's date: {datetime.date.today().strftime('%B %d, %Y')}
    
    IMPORTANT: If you have memory of writing test plans for similar features
    in past runs, reuse proven patterns and improve upon them.""",
    expected_output="""A complete, professional test plan document in markdown
    following all 12 sections of the template.""",
    agent=test_plan_writer,
    context=[analysis_task],
)


# ══════════════════════════════════════════════════════════════════
#  CREW — WITH ALL 3 MEMORY TYPES ENABLED
# ══════════════════════════════════════════════════════════════════
#
#  memory=True enables ALL three memory types at once:
#
#  ┌──────────────────────────────────────────────────────────────┐
#  │                                                              │
#  │  1️⃣  SHORT-TERM MEMORY                                      │
#  │     Scope   : Current execution only                        │
#  │     Purpose : Shares context between agents in this run     │
#  │     Example : Analyst's Jira findings → Writer uses them    │
#  │     Storage : In-memory (RAM), cleared after execution      │
#  │                                                              │
#  │  2️⃣  LONG-TERM MEMORY                                       │
#  │     Scope   : Persists across multiple executions           │
#  │     Purpose : Agents learn and improve over time            │
#  │     Example : "Last time I wrote a test plan for a cart     │
#  │               bug, I included edge cases for currency       │
#  │               rounding — I should do that again"            │
#  │     Storage : SQLite database (local, in working directory) │
#  │                                                              │
#  │  3️⃣  ENTITY MEMORY                                          │
#  │     Scope   : Persists across executions                    │
#  │     Purpose : Tracks entities (people, modules, tickets)    │
#  │     Example : "VWO-48 is about Shopping Cart. Reporter is   │
#  │               Pramod. Shopping Cart module has had 3 bugs"  │
#  │     Storage : SQLite database with entity relationships     │
#  │                                                              │
#  └──────────────────────────────────────────────────────────────┘

crew = Crew(
    agents=[ticket_analyst, test_plan_writer],
    tasks=[analysis_task, test_plan_task],
    process=Process.sequential,
    verbose=True,
    max_rpm=4,

    # ─── MEMORY CONFIGURATION ───────────────────────────────────
    memory=True,           # 🧠 Enables short-term + long-term + entity memory

    # Optional: configure the embedder for memory's vector store.
    # By default CrewAI uses a local embedder, but you can point
    # it at OpenAI, HuggingFace, etc. We use the default here.
    # embedder={
    #     "provider": "openai",
    #     "config": {"model": "text-embedding-3-small"}
    # }
)

print(f"\n🚀 Starting Test Plan Generation for {JIRA_TICKET_ID}")
print(f"🧠 Memory enabled: Short-Term + Long-Term + Entity")
print("=" * 60)

result = crew.kickoff()

analysis_output = str(analysis_task.output) if analysis_task.output else ""
test_plan_output = str(result)

print("\n" + "=" * 60)
print("📋 TEST PLAN GENERATED (with Memory)")
print("=" * 60)
print(test_plan_output[:500] + "...")


# ══════════════════════════════════════════════════════════════════
#  SAVE AS .DOCX
# ══════════════════════════════════════════════════════════════════

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("\n⚠️  python-docx not installed. Installing now...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT


def markdown_to_docx(md_text: str, ticket_id: str) -> str:
    """Convert the markdown test plan to a professional .docx document."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title page
    doc.add_paragraph()
    title = doc.add_heading(f'Test Plan — {ticket_id}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        run.font.size = Pt(28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f'Generated by AI Test Plan Agent (with Memory)\n'
        f'{datetime.date.today().strftime("%B %d, %Y")}'
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Memory badge
    mem_para = doc.add_paragraph()
    mem_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mem_para.add_run('🧠 Memory-Enhanced: Short-Term + Long-Term + Entity')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x8b, 0x5c, 0xf6)
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # Parse markdown
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Headers
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        elif stripped.startswith('## '):
            h = doc.add_heading(stripped[3:], level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x0d, 0x3b, 0x8f)
        elif stripped.startswith('# '):
            h = doc.add_heading(stripped[2:], level=0)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x0d, 0x3b, 0x8f)

        # Tables
        elif stripped.startswith('|') and stripped.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1

            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                data_start = 1
                if len(table_lines) > 1 and all(
                    c.strip().replace('-', '').replace(':', '') == ''
                    for c in table_lines[1].split('|')[1:-1]
                ):
                    data_start = 2

                data_rows = []
                for tl in table_lines[data_start:]:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    data_rows.append(cells)

                num_cols = len(headers)
                table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                header_row = table.rows[0]
                for j, h_text in enumerate(headers):
                    if j < num_cols:
                        cell = header_row.cells[j]
                        cell.text = h_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        from docx.oxml.ns import qn
                        shading = cell._element.get_or_add_tcPr()
                        shading_elm = shading.makeelement(qn('w:shd'), {
                            qn('w:fill'): '1a56db',
                            qn('w:val'): 'clear'
                        })
                        shading.append(shading_elm)

                for row_idx, row_data in enumerate(data_rows):
                    row = table.rows[row_idx + 1]
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = row.cells[col_idx]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                doc.add_paragraph()

        # Horizontal rules
        elif stripped.startswith('---') or stripped.startswith('___'):
            doc.add_paragraph('─' * 60)

        # Bullets
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # Numbered items
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # Paragraphs
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1

    # Footer
    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f'Generated by AI Test Plan Agent (Memory-Enhanced)\n'
        f'CrewAI × Groq × Jira\n'
        f'{datetime.datetime.now().strftime("%B %d, %Y at %I:%M %p")}'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)
    filename = f"TestPlan_{ticket_id}_Memory_{datetime.date.today().strftime('%Y%m%d')}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))
    return str(filepath)


# ── Generate DOCX ─────────────────────────────────────────────────
print("\n📝 Converting test plan to DOCX ...")
docx_path = markdown_to_docx(test_plan_output, JIRA_TICKET_ID)
print(f"✅ Test Plan saved to: {docx_path}")

# ── Open the document ─────────────────────────────────────────────
webbrowser.open(f"file://{Path(docx_path).resolve()}")
print("🌐 Opening document ...")
print("\n🎉 Done! Memory-enhanced test plan generation complete.")
print("🧠 Memory has been updated — next run will benefit from this execution's learnings.")